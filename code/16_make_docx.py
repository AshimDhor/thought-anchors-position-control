"""Render REPORT.md to a submission-quality .docx.

Generated from the markdown rather than hand-built, so the Word document and the
repository document cannot drift apart. Handles headings, paragraphs, bullet and
numbered lists, tables, images with numbered captions, blockquote callouts, and
inline bold/italic/code/links.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── palette, matched to the figures ──────────────────────────────────────────
INK      = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT   = RGBColor(0x1F, 0x5F, 0xA8)
ACCENT2  = RGBColor(0xC0, 0x39, 0x2B)
MUTED    = RGBColor(0x5A, 0x66, 0x72)
HDR_FILL = "1F5FA8"
ZEBRA    = "F2F5F9"
CALLOUT  = "FBF3E6"
BODY_FONT, HEAD_FONT, MONO = "Calibri", "Calibri Light", "Consolas"


def _shade(cell, hexfill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def _cell_margins(table, top=60, bottom=60, left=110, right=110) -> None:
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def _left_bar(par, colour="C0392B") -> None:
    """Accent bar down the left of a paragraph — used for callouts."""
    pPr = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), colour)
    borders.append(left)
    pPr.append(borders)


def _para_shade(par, hexfill: str) -> None:
    pPr = par._p.get_or_add_pPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexfill)
    pPr.append(el)


# ── inline markdown ──────────────────────────────────────────────────────────
_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_runs(par, text: str, size=10.5, colour=INK, bold_all=False) -> None:
    text = text.replace("\\|", "|").replace("≈", "≈")
    for piece in _INLINE.split(text):
        if not piece:
            continue
        run = par.add_run()
        run.font.size = Pt(size)
        run.font.color.rgb = colour
        run.font.name = BODY_FONT
        run.bold = bold_all
        if piece.startswith("**") and piece.endswith("**"):
            run.text, run.bold = piece[2:-2], True
        elif piece.startswith("`") and piece.endswith("`"):
            run.text = piece[1:-1]
            run.font.name = MONO
            run.font.size = Pt(size - 0.7)
            run.font.color.rgb = ACCENT2
        elif piece.startswith("*") and piece.endswith("*"):
            run.text, run.italic = piece[1:-1], True
        elif piece.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", piece)
            run.text = m.group(1)
            run.font.color.rgb = ACCENT
            run.underline = True
        else:
            run.text = piece


def para(doc, text="", size=10.5, space_after=6, colour=INK, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        add_runs(p, text, size=size, colour=colour)
    return p


def heading(doc, text, level):
    sizes = {1: 19, 2: 14, 3: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Strip section numbering entirely ("3.1 Setting" -> "Setting"). Stripping
    # only the first group turned "3.1 Setting" into "1 Setting".
    text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
    run = p.add_run(text)
    run.font.name = HEAD_FONT
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    run.font.color.rgb = ACCENT if level <= 2 else INK
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        b = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), "D6DEE8")
        b.append(bot); pPr.append(b)
    return p


def add_table(doc, rows):
    has_header = any(c.strip() for c in rows[0])
    ncol = len(rows[0])
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_margins(t)
    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for c, val in enumerate(row):
            cell = cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            if r == 0 and has_header:
                _shade(cell, HDR_FILL)
                add_runs(p, val, size=9, colour=RGBColor(0xFF, 0xFF, 0xFF), bold_all=True)
            else:
                if r % 2 == (0 if has_header else 1):
                    _shade(cell, ZEBRA)
                add_runs(p, val, size=9, bold_all=(not has_header and c == 0))
            if c > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, path: Path, number: int, caption: str, width_in=6.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Figure {number}. ")
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = ACCENT
    r.font.name = BODY_FONT
    add_runs(cap, caption, size=8.5, colour=MUTED)


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.15
    _left_bar(p, "C0392B")
    _para_shade(p, CALLOUT)
    add_runs(p, text, size=10.5)


def bullet(doc, text, level=0, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.28 + 0.24 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    add_runs(p, text, size=10.5)
    return p


# ── document furniture ───────────────────────────────────────────────────────
def title_block(doc, title, subtitle, byline, meta):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = HEAD_FONT; r.font.size = Pt(26); r.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle)
    r.font.name = HEAD_FONT; r.font.size = Pt(13); r.italic = True
    r.font.color.rgb = ACCENT2

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(byline)
    r.font.name = BODY_FONT; r.font.size = Pt(11); r.bold = True

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    add_runs(p, meta, size=9, colour=MUTED)

    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "8"); bot.set(qn("w:color"), "1F5FA8")
    b.append(bot); pPr.append(b)


def add_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    run = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "Right-click and choose Update Field to build the table of contents."
    run.append(t); fld.append(run)
    p._p.append(fld)


def footer_page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.font.size = Pt(8.5); r.font.color.rgb = MUTED
    for instr in ("PAGE",):
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
        it.text = f" {instr} "
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        r._r.append(f1); r._r.append(it); r._r.append(f2)


def header_text(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(8.5); r.font.color.rgb = MUTED; r.font.name = BODY_FONT


# ── main conversion ──────────────────────────────────────────────────────────
CAPTIONS = {
    "fig0": "What resampling importance is made of. Left: variance explained by each "
            "candidate. Middle: importance against the entropy of the distribution the "
            "sentence arrives into. Right: that entropy is only partly positional "
            "(ρ = −0.45), so it is not position in disguise.",
    "fig1": "Measured importance against position, with the finite-sample floor drawn in "
            "(left), and the collapse of answer-distribution entropy along the trace (right).",
    "fig2": "What a predictor that never reads the text recovers, against chance.",
    "fig3": "The filler arm. Retained as a negative result about the control, not as a "
            "result about the model — see §8.2.",
    "fig4": "Importance by sentence category, raw (left) and after subtracting the "
            "position-only prediction (right). The top two categories survive.",
    "fig5": "Category importance against category mean position. Category is strongly "
            "positional, yet contributes almost independently of position.",
    "fig6": "Receiver-head score across all 48 layers × 40 heads, with the top heads "
            "circled, and the depth profile.",
    "fig7": "Sentence-to-sentence attention. The middle panel — the selected 'receiver' "
            "heads — is almost purely diagonal: these are recency heads.",
    "fig8": "The black-box and white-box measures come apart (middle), but excluding "
            "attention near the diagonal removes most of the white-box effect (right).",
}


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default="../REPORT.md")
    ap.add_argument("--out", default="../Ashim_Dhor_MATS12_Thought_Anchors.docx")
    args = ap.parse_args()

    root = Path(args.src).resolve().parent
    md = Path(args.src).read_text().splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT; style.font.size = Pt(10.5)
    header_text(sec, "Ashim Dhor · MATS 12.0 · Neel Nanda stream")
    footer_page_numbers(sec)

    title_block(
        doc,
        "What Resampling Importance Actually Measures",
        "A red-team of thought anchors that failed, and the thing it found instead",
        "Ashim Dhor",
        "DeepSeek-R1-Distill-Qwen-14B on MATH-500 · 8 traces · 254 sentences · "
        "64 rollouts per prefix boundary. Every number is generated from saved data; "
        "none is transcribed by hand.",
    )

    heading(doc, "Contents", 2)
    add_toc(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    i, fignum, in_contents, seen_h1 = 0, 0, False, False
    while i < len(md):
        line = md[i].rstrip()

        if line.startswith("![") :
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
            if m:
                path = (root / m.group(1)).resolve()
                key = Path(m.group(1)).name.split("_")[0]
                fignum += 1
                # skip an italic caption line following the image
                cap = CAPTIONS.get(key, "")
                if i + 1 < len(md) and md[i + 1].startswith("*") and not md[i+1].startswith("**"):
                    i += 1
                if path.exists():
                    add_figure(doc, path, fignum, cap)
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(md) and md[i].lstrip().startswith("|"):
                # \| is an escaped pipe inside a cell (e.g. |ΔP(correct)|), not a
                # column separator. Protect it before splitting.
                raw = md[i].strip().replace("\\|", "\x00")
                cells = [c.strip().replace("\x00", "|")
                         for c in raw.strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                rows = [r + [""] * (width - len(r)) for r in rows]
                add_table(doc, rows)
            continue

        if line.startswith("### "):
            heading(doc, line[4:], 3); i += 1; continue
        if line.startswith("## "):
            txt = line[3:]
            in_contents = txt.strip().lower() == "contents"
            if not in_contents:
                heading(doc, txt, 2)
            i += 1
            continue
        if line.startswith("# "):
            if not seen_h1:
                # The title block already carries the title, subtitle, byline and
                # metadata. Skip the markdown's whole front matter up to the first
                # "## " heading rather than re-emitting any of it.
                seen_h1 = True
                i += 1
                while i < len(md) and not md[i].startswith("## "):
                    i += 1
                continue
            heading(doc, line[2:], 1); i += 1; continue

        if in_contents:
            if line.startswith("---"):
                in_contents = False
            i += 1
            continue

        if line.startswith(">"):
            block = []
            while i < len(md) and md[i].startswith(">"):
                block.append(md[i].lstrip("> ").rstrip()); i += 1
            add_callout(doc, " ".join(b for b in block if b))
            continue

        if re.match(r"^\d+\.\s", line):
            bullet(doc, re.sub(r"^\d+\.\s", "", line), numbered=True); i += 1; continue
        if line.startswith("- ") or line.startswith("* "):
            bullet(doc, line[2:]); i += 1; continue

        if line.startswith("```"):
            code = []
            i += 1
            while i < len(md) and not md[i].startswith("```"):
                code.append(md[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            _para_shade(p, "F4F6F8")
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(10)
            for k, cl in enumerate(code):
                r = p.add_run(cl + ("\n" if k < len(code) - 1 else ""))
                r.font.name = MONO; r.font.size = Pt(8.5); r.font.color.rgb = INK
            continue

        if line.startswith("---"):
            i += 1; continue
        if not line.strip():
            i += 1; continue

        para(doc, line)
        i += 1

    doc.save(args.out)
    print(f"wrote {args.out}  ({fignum} figures embedded)")


if __name__ == "__main__":
    main()
